#!/usr/bin/env python3
from __future__ import annotations

import argparse
import csv
import json
import sys
import zipfile
from pathlib import Path
from typing import Iterable

from docx import Document
from openpyxl import load_workbook
from odf import opendocument, table, text

TEXT_EXTS = {".docx", ".odt", ".csv", ".tsv"}
SHEET_EXTS = {".xlsx", ".xlsm", ".ods"}
ZIP_MEDIA_EXTS = {".docx", ".xlsx", ".xlsm", ".odt", ".ods"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".svg"}


def ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def clean_text(value: object) -> str:
    if value is None:
        return ""
    text_value = str(value).replace("\r\n", "\n").replace("\r", "\n").strip()
    return text_value


def markdown_escape_cell(value: object) -> str:
    text_value = clean_text(value)
    if not text_value:
        return ""
    return text_value.replace("|", "\\|").replace("\n", "<br>")


def write_output(content: str, output: Path | None) -> None:
    if output is None:
        sys.stdout.write(content)
        if content and not content.endswith("\n"):
            sys.stdout.write("\n")
        return
    ensure_parent(output)
    output.write_text(content, encoding="utf-8")


def describe_type(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".docx", ".odt"}:
        return "document"
    if suffix in {".xlsx", ".xlsm", ".ods", ".csv", ".tsv"}:
        return "spreadsheet"
    return "unknown"


def count_zip_media(path: Path) -> int:
    if path.suffix.lower() not in ZIP_MEDIA_EXTS:
        return 0
    try:
        with zipfile.ZipFile(path) as zf:
            return sum(1 for name in zf.namelist() if Path(name).suffix.lower() in IMAGE_SUFFIXES)
    except Exception:
        return 0


def inspect_docx(path: Path) -> dict:
    doc = Document(path)
    paragraphs = [clean_text(p.text) for p in doc.paragraphs if clean_text(p.text)]
    return {
        "path": str(path),
        "format": ".docx",
        "kind": "document",
        "paragraphs": len(paragraphs),
        "tables": len(doc.tables),
        "media_files": count_zip_media(path),
        "preview": paragraphs[:8],
    }


def render_docx_markdown(path: Path) -> str:
    doc = Document(path)
    lines = [f"# DOCX: {path.name}", ""]
    body_started = False
    for paragraph in doc.paragraphs:
        text_value = clean_text(paragraph.text)
        if not text_value:
            continue
        style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
        if style_name.startswith("Heading"):
            level_text = "".join(ch for ch in style_name if ch.isdigit())
            level = int(level_text) if level_text else 2
            level = max(1, min(level, 6))
            lines.append(f"{'#' * level} {text_value}")
        else:
            lines.append(text_value)
        lines.append("")
        body_started = True

    for index, tbl in enumerate(doc.tables, 1):
        rows = []
        for row in tbl.rows:
            values = [markdown_escape_cell(cell.text) for cell in row.cells]
            if any(values):
                rows.append(values)
        if not rows:
            continue
        if not body_started:
            lines.append("")
        lines.append(f"## Table {index}")
        width = max(len(r) for r in rows)
        normalized = [r + [""] * (width - len(r)) for r in rows]
        header = normalized[0]
        divider = ["---"] * width
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join(divider) + " |")
        for row in normalized[1:]:
            lines.append("| " + " | ".join(row) + " |")
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def inspect_xlsx(path: Path) -> dict:
    wb = load_workbook(path, read_only=True, data_only=True)
    sheets = []
    for ws in wb.worksheets:
        preview = []
        for row in ws.iter_rows(min_row=1, max_row=5, values_only=True):
            values = [clean_text(v) for v in row]
            if any(values):
                preview.append(values)
        sheets.append({
            "name": ws.title,
            "max_row": ws.max_row,
            "max_column": ws.max_column,
            "preview": preview,
        })
    return {
        "path": str(path),
        "format": path.suffix.lower(),
        "kind": "spreadsheet",
        "sheet_count": len(sheets),
        "sheets": sheets,
        "media_files": count_zip_media(path),
    }


def render_xlsx_markdown(path: Path, row_limit: int = 20) -> str:
    wb = load_workbook(path, read_only=True, data_only=True)
    lines = [f"# Workbook: {path.name}", ""]
    for ws in wb.worksheets:
        lines.append(f"## Sheet: {ws.title}")
        lines.append(f"- Rows: {ws.max_row}")
        lines.append(f"- Columns: {ws.max_column}")
        lines.append("")
        preview_rows = []
        for row in ws.iter_rows(min_row=1, max_row=row_limit, values_only=True):
            values = [markdown_escape_cell(v) for v in row]
            if any(values):
                preview_rows.append(values)
        if preview_rows:
            width = max(len(r) for r in preview_rows)
            normalized = [r + [""] * (width - len(r)) for r in preview_rows]
            header = normalized[0]
            divider = ["---"] * width
            lines.append("| " + " | ".join(header) + " |")
            lines.append("| " + " | ".join(divider) + " |")
            for row in normalized[1:]:
                lines.append("| " + " | ".join(row) + " |")
        else:
            lines.append("[No non-empty rows in preview]")
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def inspect_delimited(path: Path, delimiter: str) -> dict:
    rows = []
    with path.open("r", encoding="utf-8", newline="") as handle:
        reader = csv.reader(handle, delimiter=delimiter)
        for idx, row in enumerate(reader, 1):
            rows.append(row)
            if idx >= 5:
                break
    total_rows = sum(1 for _ in path.open("r", encoding="utf-8", newline=""))
    return {
        "path": str(path),
        "format": path.suffix.lower(),
        "kind": "spreadsheet",
        "rows": total_rows,
        "preview": rows,
    }


def render_delimited_markdown(path: Path, delimiter: str, row_limit: int = 20) -> str:
    rows = []
    with path.open("r", encoding="utf-8", newline="") as handle:
        reader = csv.reader(handle, delimiter=delimiter)
        for idx, row in enumerate(reader, 1):
            values = [markdown_escape_cell(v) for v in row]
            if any(values):
                rows.append(values)
            if idx >= row_limit:
                break
    lines = [f"# Table: {path.name}", ""]
    if not rows:
        lines.append("[No non-empty rows]")
        return "\n".join(lines) + "\n"
    width = max(len(r) for r in rows)
    normalized = [r + [""] * (width - len(r)) for r in rows]
    header = normalized[0]
    divider = ["---"] * width
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(divider) + " |")
    for row in normalized[1:]:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")
    return "\n".join(lines)


def node_text(node) -> str:
    parts = []
    for child in getattr(node, 'childNodes', []):
        if getattr(child, 'nodeType', None) == 3:
            parts.append(str(child.data))
        else:
            parts.append(node_text(child))
    return ''.join(parts)


def inspect_odt(path: Path) -> dict:
    doc = opendocument.load(str(path))
    paragraphs = [clean_text(node_text(p)) for p in doc.getElementsByType(text.P)]
    paragraphs = [p for p in paragraphs if p]
    return {
        "path": str(path),
        "format": ".odt",
        "kind": "document",
        "paragraphs": len(paragraphs),
        "media_files": count_zip_media(path),
        "preview": paragraphs[:8],
    }


def render_odt_markdown(path: Path) -> str:
    doc = opendocument.load(str(path))
    lines = [f"# ODT: {path.name}", ""]
    for heading in doc.getElementsByType(text.H):
        text_value = clean_text(node_text(heading))
        if text_value:
            lines.append(f"## {text_value}")
            lines.append("")
    for paragraph in doc.getElementsByType(text.P):
        text_value = clean_text(node_text(paragraph))
        if text_value:
            lines.append(text_value)
            lines.append("")
    return "\n".join(lines).strip() + "\n"


def inspect_ods(path: Path) -> dict:
    doc = opendocument.load(str(path))
    sheet_infos = []
    for sheet in doc.spreadsheet.getElementsByType(table.Table):
        rows = sheet.getElementsByType(table.TableRow)
        sheet_infos.append({
            "name": str(sheet.getAttribute("name") or "Sheet"),
            "rows": len(rows),
        })
    return {
        "path": str(path),
        "format": ".ods",
        "kind": "spreadsheet",
        "sheet_count": len(sheet_infos),
        "sheets": sheet_infos,
        "media_files": count_zip_media(path),
    }


def iter_ods_rows(sheet, row_limit: int) -> Iterable[list[str]]:
    count = 0
    for row in sheet.getElementsByType(table.TableRow):
        values = []
        for cell in row.getElementsByType(table.TableCell):
            repeat = int(cell.getAttribute("numbercolumnsrepeated") or 1)
            value = clean_text(node_text(cell))
            for _ in range(repeat):
                values.append(value)
        if any(values):
            yield values
            count += 1
            if count >= row_limit:
                break


def render_ods_markdown(path: Path, row_limit: int = 20) -> str:
    doc = opendocument.load(str(path))
    lines = [f"# ODS: {path.name}", ""]
    for sheet in doc.spreadsheet.getElementsByType(table.Table):
        name = str(sheet.getAttribute("name") or "Sheet")
        lines.append(f"## Sheet: {name}")
        rows = list(iter_ods_rows(sheet, row_limit))
        if not rows:
            lines.append("[No non-empty rows in preview]")
            lines.append("")
            continue
        width = max(len(r) for r in rows)
        normalized = [[markdown_escape_cell(v) for v in r] + [""] * (width - len(r)) for r in rows]
        header = normalized[0]
        divider = ["---"] * width
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join(divider) + " |")
        for row in normalized[1:]:
            lines.append("| " + " | ".join(row) + " |")
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def inspect_file(path: Path) -> dict:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return inspect_docx(path)
    if suffix in {".xlsx", ".xlsm"}:
        return inspect_xlsx(path)
    if suffix == ".csv":
        return inspect_delimited(path, ",")
    if suffix == ".tsv":
        return inspect_delimited(path, "\t")
    if suffix == ".odt":
        return inspect_odt(path)
    if suffix == ".ods":
        return inspect_ods(path)
    raise SystemExit(f"Unsupported file type: {suffix}")


def extract_to_markdown(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return render_docx_markdown(path)
    if suffix in {".xlsx", ".xlsm"}:
        return render_xlsx_markdown(path)
    if suffix == ".csv":
        return render_delimited_markdown(path, ",")
    if suffix == ".tsv":
        return render_delimited_markdown(path, "\t")
    if suffix == ".odt":
        return render_odt_markdown(path)
    if suffix == ".ods":
        return render_ods_markdown(path)
    raise SystemExit(f"Unsupported file type: {suffix}")


def export_xlsx_sheets(path: Path, output_dir: Path) -> list[str]:
    output_dir.mkdir(parents=True, exist_ok=True)
    wb = load_workbook(path, read_only=True, data_only=True)
    outputs = []
    for ws in wb.worksheets:
        safe_name = ''.join(ch if ch.isalnum() or ch in ('-', '_') else '_' for ch in ws.title).strip('_') or 'sheet'
        target = output_dir / f"{safe_name}.csv"
        with target.open('w', encoding='utf-8', newline='') as handle:
            writer = csv.writer(handle)
            for row in ws.iter_rows(values_only=True):
                writer.writerow(['' if v is None else v for v in row])
        outputs.append(str(target))
    return outputs


def export_ods_sheets(path: Path, output_dir: Path) -> list[str]:
    output_dir.mkdir(parents=True, exist_ok=True)
    doc = opendocument.load(str(path))
    outputs = []
    for sheet in doc.spreadsheet.getElementsByType(table.Table):
        name = str(sheet.getAttribute("name") or "sheet")
        safe_name = ''.join(ch if ch.isalnum() or ch in ('-', '_') else '_' for ch in name).strip('_') or 'sheet'
        target = output_dir / f"{safe_name}.csv"
        with target.open('w', encoding='utf-8', newline='') as handle:
            writer = csv.writer(handle)
            for row in sheet.getElementsByType(table.TableRow):
                values = []
                for cell in row.getElementsByType(table.TableCell):
                    repeat = int(cell.getAttribute("numbercolumnsrepeated") or 1)
                    value = clean_text(node_text(cell))
                    for _ in range(repeat):
                        values.append(value)
                writer.writerow(values)
        outputs.append(str(target))
    return outputs


def extract_media(path: Path, output_dir: Path) -> list[str]:
    if path.suffix.lower() not in ZIP_MEDIA_EXTS:
        raise SystemExit(f"Media extraction is not supported for: {path.suffix.lower()}")
    output_dir.mkdir(parents=True, exist_ok=True)
    written = []
    with zipfile.ZipFile(path) as zf:
        for name in zf.namelist():
            member = Path(name)
            suffix = member.suffix.lower()
            if suffix not in IMAGE_SUFFIXES:
                continue
            lower_name = name.lower()
            if ('media/' not in lower_name) and (not lower_name.startswith('pictures/')):
                continue
            target = output_dir / member.name
            counter = 1
            while target.exists():
                target = output_dir / f"{member.stem}_{counter}{member.suffix}"
                counter += 1
            with zf.open(name) as src, target.open('wb') as dst:
                dst.write(src.read())
            written.append(str(target))
    return written


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description='Inspect and extract Office/OpenDocument files')
    sub = parser.add_subparsers(dest='cmd', required=True)

    inspect_p = sub.add_parser('inspect', help='Print JSON summary for an office file')
    inspect_p.add_argument('--input', required=True, type=Path)

    extract_p = sub.add_parser('extract', help='Extract readable markdown from an office file')
    extract_p.add_argument('--input', required=True, type=Path)
    extract_p.add_argument('--output', type=Path)

    export_p = sub.add_parser('export-sheets', help='Export spreadsheet sheets to CSV files')
    export_p.add_argument('--input', required=True, type=Path)
    export_p.add_argument('--output-dir', required=True, type=Path)

    media_p = sub.add_parser('extract-media', help='Extract embedded images from Office zip-based files')
    media_p.add_argument('--input', required=True, type=Path)
    media_p.add_argument('--output-dir', required=True, type=Path)

    return parser.parse_args()


def main() -> int:
    args = parse_args()
    input_path: Path = args.input.expanduser().resolve()
    if not input_path.exists():
        raise SystemExit(f'Input file does not exist: {input_path}')

    if args.cmd == 'inspect':
        summary = inspect_file(input_path)
        print(json.dumps(summary, ensure_ascii=False, indent=2))
        return 0

    if args.cmd == 'extract':
        content = extract_to_markdown(input_path)
        output = args.output.expanduser() if args.output else None
        write_output(content, output)
        return 0

    if args.cmd == 'export-sheets':
        output_dir = args.output_dir.expanduser()
        suffix = input_path.suffix.lower()
        if suffix in {'.xlsx', '.xlsm'}:
            outputs = export_xlsx_sheets(input_path, output_dir)
        elif suffix == '.ods':
            outputs = export_ods_sheets(input_path, output_dir)
        else:
            raise SystemExit(f'export-sheets only supports .xlsx, .xlsm, and .ods (got {suffix})')
        print(json.dumps({"written": outputs}, ensure_ascii=False, indent=2))
        return 0

    if args.cmd == 'extract-media':
        output_dir = args.output_dir.expanduser()
        outputs = extract_media(input_path, output_dir)
        print(json.dumps({"written": outputs}, ensure_ascii=False, indent=2))
        return 0

    raise SystemExit('Unknown command')


if __name__ == '__main__':
    raise SystemExit(main())
